"""Markdown to DOCX converter using python-docx."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_run_with_style(para, text, bold=False, italic=False, code=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    return run


def parse_inline(para, text):
    """Parse inline markdown: **bold**, `code`, plain text."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            add_run_with_style(para, text[last:m.start()])
        if m.group(2):  # bold
            add_run_with_style(para, m.group(2), bold=True)
        elif m.group(3):  # code
            add_run_with_style(para, m.group(3), code=True)
        last = m.end()
    if last < len(text):
        add_run_with_style(para, text[last:])


def add_image(doc, img_path, caption_text, max_width_cm=14):
    from docx.shared import Cm
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(max_width_cm))
    except Exception as e:
        p = doc.add_paragraph(f'[그림: {img_path.name}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_text:
        cap = doc.add_paragraph(caption_text)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True


def add_table(doc, rows):
    """rows: list of lists of strings. First row = header."""
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_bg(cell, 'D9D9D9')
    doc.add_paragraph()


def convert(md_path, out_path):
    md_path = Path(md_path)
    out_path = Path(out_path)
    base_dir = md_path.parent

    lines = md_path.read_text(encoding='utf-8').splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # Blank line
        if line.strip() == '':
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Image
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            caption = img_match.group(1)
            rel_path = img_match.group(2)
            img_path = base_dir / rel_path
            add_image(doc, img_path, caption)
            i += 1
            continue

        # Table
        if line.strip().startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                row_line = lines[i].strip().strip('|')
                if re.match(r'^[\s|:-]+$', row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split('|')]
                table_lines.append(cells)
                i += 1
            if table_lines:
                add_table(doc, table_lines)
            continue

        # Code block
        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            # Light grey background via paragraph shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)
            continue

        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < n and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('> ').rstrip())
                i += 1
            # Merge and add as indented paragraph
            merged = ' '.join(quote_lines)
            # Handle nested code inside blockquote (simple: just strip)
            merged = re.sub(r'```[a-z]*', '', merged)
            p = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            parse_inline(p, merged)
            continue

        # Bullet list (- or *)
        if re.match(r'^(\s*)[-*] ', line):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            text = re.sub(r'^(\s*)[-*] ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
            parse_inline(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, text)
            i += 1
            continue

        # Checkbox list
        if re.match(r'^- \[[ x]\] ', line):
            text = re.sub(r'^- \[[ x]\] ', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            parse_inline(p, '☐ ' + text)
            i += 1
            continue

        # Normal paragraph
        para_lines = []
        while i < n and lines[i].strip() != '' and not lines[i].startswith('#') \
                and not lines[i].startswith('>') and not lines[i].startswith('```') \
                and not lines[i].startswith('|') and not re.match(r'^---+$', lines[i].strip()) \
                and not lines[i].startswith('!'):
            para_lines.append(lines[i].rstrip())
            i += 1
        text = ' '.join(para_lines)
        if text.strip():
            p = doc.add_paragraph()
            parse_inline(p, text)
        continue

    doc.save(str(out_path))
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    md = sys.argv[1]
    out = sys.argv[2]
    convert(md, out)
